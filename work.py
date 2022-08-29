import docx
from random import randint
from num2t4ru import decimal2text
from datetime import datetime

class do_work(object):
	def __init__(self):
		super(self).__init__()
		self.doc = docx.Document('default.docx')
		self.fontstyles = self.doc.styles
		self.user = 2
		self.int_units  = ((u'рубль', u'рубля', u'рублей'), 'm')
		self.exp_units  = ((u'копейка', u'копейки', u'копеек'), 'f')
		self.customer = {'company':'', 'comdir':'', 'compdoc':'', 'comaddr':'', 'combank':'', 'comunp':'', 'comletter':'', 'theme':'', 'themeadd': '', 'allmoney' : ''}
		self.pay = {'payme' : '', 'stealsoc' : '', 'stealins' : '', 'stealnakl' : '', 'payandsteal' : '', 'profit' : '', 'withoutndssteal' : '', 'ndssteal' : '', 'allmoney' : ''}
		self.w_pay = {'dgwithoutndssteal': '', 'dgndssteal': '', 'dgallmoney': '', 'wwithoutndssteal': '', 'wndssteal': '', 'wallmoney': ''}
		self.decode  = [[], [], [], [], []]
		self.timestamp = {'year': '', 'month': '','monthtill': ''}
		self.months = {1: 'января', 2: 'февраля',3: 'марта',4: 'апреля',5: 'мая',6: 'июня',7: 'июля',8: 'августа',9: 'сентября',10: 'октября',11: 'ноября',12: 'декабря'}
		self.bntu = {}
		self.pay_calc = {}
		self.coeff = {'техник': 12,'инж.': 15,'инж.1к': 17,'инж.2к': 19,'вед.инж.': 25,'зав.сект.': 25,'м.н.с.': 25,'н.с.': 27,'с.н.с.': 30,'в.н.с.': 35, 'зав.лаб.': 10}

		self.timestamp['year'] = list(datetime.now().timetuple())[0]%2000
		self.timestamp['month'] = self.months[list(datetime.now().timetuple())[1]]
		if list(datetime.now().timetuple())[1] + 6 > 12:
			self.timestamp['monthtill'] = self.months[12]
		else:
			self.timestamp['monthtill'] = self.months[list(datetime.now().timetuple())[1]+6]
		with open('work.txt', 'r') as in_values:
			n = 0
			for line in in_values:
				if n < len(self.customer):
					self.customer[list(self.customer.keys())[n]] = line[:-1]
					n+=1
		with open('users.txt', 'r') as bntu_val:
			for line in bntu_val:
				self.bntu[line.split('%')[0]] = line.split('%')[1][:-1]

		tab_calc = 16
		for row in self.doc.tables[tab_calc].rows[2:]:
			for j in [-1, -2]:
				if len(row.cells[j].text)!= 0:
					self.pay_calc[row.cells[j].text] = ''

		self.decode_pay()
		self.count_money()
		self.calculation_pay()
		self.num_to_text()
		self.calculation_fill()
		#big_Dict
		final = {}
		final.update(customer)
		final.update(pay)
		final.update(w_pay)
		final.update(timestamp)
		final.update(bntu)
		final.update(pay_calc)
		for i in range(len(self.doc.paragraphs)):
			fill_text(self.doc.paragraphs[i])
		self.read_calc()
		self.file_name = 'EKT'
		self.doc.save(file_name + '.docx')


	#second_step count money
	def count_money(self):
		self.pay['allmoney'] = round(float(self.customer['allmoney']), 2)
		self.pay['ndssteal'] = round(float(self.pay['allmoney'])/6, 2)
		self.pay['withoutndssteal']= round(self.pay['allmoney'] - self.pay['ndssteal'], 2)
		self.pay['profit'] = round(float(self.pay['withoutndssteal'])/11, 2)
		self.pay['payandsteal'] = round(float(self.pay['withoutndssteal'])*10/11, 2)
		self.pay['payme'] = round(float(self.pay['payandsteal'])/1.691, 2)
		self.pay['stealins'] = round(float(self.pay['payme'])*0.001, 2)
		self.pay['stealsoc'] = round(float(self.pay['payme'])*0.34, 2)
		self.pay['stealnakl'] = round(self.pay['payandsteal'] - self.pay['stealins'] - self.pay['stealsoc'] - self.pay['payme'], 2)

	#third step calculate cost
	def calculation_pay(self):
		for i in range(len(self.decode [0])):
			self.decode[3][i] = round(self.pay['payme'] * self.coeff[self.decode[0][i]]/110.2 * 171.2 * self.decode[1][i]/self.decode[2][i])
			self.decode[4][i] = round(self.decode[3][i] * self.decode[2][i] / (171.2 * self.decode[1][i]), 2)
		self.decode[4][user] = round(self.decode[4][user] + pay['payme'] - sum(self.decode[4]), 2)
		self.decode[3][user] = round(self.decode[4][user]* 171.2 * self.decode[1][user] / self.decode[2][user], 2)

	#fourth step
	def calculation_fill(self):
		self.decode_buf = self.decode[3]+self.decode [4]
		self.decode_buf.append(self.pay['payme'])
		i = 0
		
		for key in sorted(list(pay_calc.keys())):
			self.pay_calc[key] = self.decode_buf[i]
			i += 1

	#fifth step cost to language
	def num_to_text(self):
		self.w_pay['dgwithoutndssteal'] = f"{int(self.pay['withoutndssteal'])} руб. {int(self.pay['withoutndssteal']*100%100)} коп.".replace(' 0 ', ' 00 ')
		self.w_pay['dgndssteal'] = f"{int(self.pay['ndssteal'])} руб. {int(self.pay['ndssteal']*100%100)} коп.".replace(' 0 ', ' 00 ')
		self.w_pay['dgallmoney'] = f"{int(self.pay['allmoney'])} руб. {int(self.pay['allmoney']*100%100)} коп.".replace(' 0 ', ' 00 ')
		self.w_pay['wwithoutndssteal'] = decimal2text(str(self.pay['withoutndssteal']),int_units = self.int_units ,exp_units =self.exp_units ).capitalize()
		self.w_pay['wndssteal'] = decimal2text(str(self.pay['ndssteal']),int_units =self.int_units , exp_units =self.exp_units ).capitalize()
		self.w_pay['wallmoney'] = decimal2text(str(self.pay['allmoney']),int_units =self.int_units , exp_units =self.exp_units ).capitalize()

	#sixth step fill paragraph
	def fill_text(self, element):
		for run in element.runs:
			if (run.text) in final.keys():
				run.text=run.text.replace(run.text, str(final[run.text]))

	#seventh step fill tables
	def read_calc(self):
		for i in range(len(self.doc.tables)):
			for j in range(len(self.doc.tables[i].rows)):
				for k in range(len(self.doc.tables[i].rows[0].cells)):
					cell = self.doc.tables[i].cell(j, k)
					for L in range(len(cell.paragraphs)):
						fill_text(cell.paragraphs[L])

	#decoding
	#tab = [post, rate, pay, hours, full]
	def decode_pay(self):
		for row in self.self.doc.tables[16].rows[2:-1]:
			#print('\''+row.cells[3].text+'\'' + ' : '+'\'\'')
			self.decode[0].append(row.cells[0].text)
			self.decode[1].append(float(row.cells[1].text.replace(',', '.')))
			self.decode[2].append(float(row.cells[-3].text.replace(',', '.')))
			self.decode[3].append(0)
			self.decode[4].append(0)


if __name__ == "__main__":
	do_work()