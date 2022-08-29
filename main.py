from PyQt5 import uic
from PyQt5.QtPrintSupport import QPrintDialog, QPrinter, QPrintPreviewDialog
from PyQt5.QtGui import QTextDocument, QTextCursor, QResizeEvent, QBrush, QColor, QTextCursor
from PyQt5.QtWidgets import QLabel, QComboBox, QApplication, QMainWindow, QDialog, QFileDialog, QMessageBox, QInputDialog,  QMainWindow, QGridLayout, QWidget, QTableWidget, QTableWidgetItem, QGraphicsScene, QLineEdit, QPushButton, QVBoxLayout
from PyQt5.QtCore import QSize, Qt, pyqtSignal, QMetaObject, QRect, QCoreApplication, QEvent
from os import listdir, path, remove
from pickle import load
import docx
from random import randint
from num2t4ru import decimal2text
from datetime import datetime

Form, _ = uic.loadUiType("interface.ui")

class Ui(QMainWindow, Form):
	resized=pyqtSignal()

	def __init__(self):
		super(Ui, self).__init__()
		self.setupUi(self)
		self.init_base()
		self.customer = {}
		self.search_1.textChanged.connect(self.change)
		self.search_2.textChanged.connect(self.change)
		self.comboBox.currentIndexChanged.connect(self.combochange)
		self.first_table = self.prices[list(self.prices.keys())[0]]
		self.addButton_1.clicked.connect(self.add_to_final)
		self.deleteButton.clicked.connect(self.delete_row)
		self.fill_users()
		self.second_table = []
		self.table_show = 0
		self.sum_raw = 0
		self.sum_full = 0
		self.koeff = 4360
		self.sale.setInputMask("99")
		#self.textfield_9.setInputMask("9999")
		#self.textfield_10.setInputMask("9999")
		self.tableWidget.horizontalHeader().setVisible(True)
		self.tableWidget_2.horizontalHeader().setVisible(True)
		self.tableWidget_2.itemChanged.connect(self.change_saving) #результаты
		self.sale.textChanged.connect(self.calculate_money)
		self.pushButton_4.clicked.connect(self.export_to_doc)
		self.pay = {'payme' : '', 'stealsoc' : '', 'stealins' : '', 'stealnakl' : '', 'payandsteal' : '', 'profit' : '', 'withoutndssteal' : '', 'ndssteal' : '', 'allmoney' : ''}
		self.w_pay = {'dgwithoutndssteal': '', 'dgndssteal': '', 'dgallmoney': '', 'wwithoutndssteal': '', 'wndssteal': '', 'wallmoney': ''}
		self.decode  = [[], [], [], [], []]
		self.timestamp = {'year': '', 'month': '','monthtill': '', 'yeartill': ''}
		self.bntu = {'nipi': '', 'nipidoc': '', 'nipidol': '', 'chef': '', 'boss': '', 'zavpes': ''}
		self.pay_calc = {}
		self.coeff = {'техник': 12,'инж.': 15,'инж.1к': 17,'инж.2к': 19,'вед.инж.': 25,'зав.сект.': 25,'м.н.с.': 25,'н.с.': 27,'с.н.с.': 30,'в.н.с.': 35, 'зав.лаб.': 10}
		self.timestamp['year'] = list(datetime.now().timetuple())[0]%2000
		self.timestamp['month'] = self.months[list(datetime.now().timetuple())[1]]
		self.fields = [self.textfield_1, self.textfield_2,self.textfield_12,self.textfield_3,self.textfield_4,self.textfield_5,self.textfield_6,self.textfield_7,self.textfield_8,self.textfield_9, self.textfield_10,self.textfield_14,self.textfield_15,self.textfield_16,self.textfield_17,self.textfield_19,self.textfield_18]
		self.openF.triggered.connect(self.Open_file)
		self.pushButton_3.clicked.connect(self.Save_file)
		self.NewDoc.triggered.connect(self.clear_window)
		self.dops = True

	def add_to_final(self):
		if self.tableWidget.selectedItems():
			for elem in self.tableWidget.selectedItems():
				self.second_table.append(list(self.first_table[elem.row()]))
				self.table_create(self.tableWidget_2, self.second_table)
				self.calculate_money()

	def clear_window(self):
		
		for i in range(len(self.fields)):
			self.fields[i].clear()
		self.tableWidget.clearContents()
		self.first_table = self.prices[list(self.prices.keys())[0]]
		self.tableWidget_2.clearContents()
		self.second_table = []
		self.sale.setText('0')

	def change_saving(self, instance): #перенос изменений в главную таблицу
		if self.table_show == 0:
			for elem in (self.tableWidget_2.selectedItems()):
				if elem.text().isdigit():
					if elem.column() == 1:
						self.second_table[elem.row()][elem.column()] = int(elem.text())
					elif elem.column() == 2:
						self.second_table[elem.row()][elem.column()] = float(elem.text())
				else:
					if elem.column() == 0:	
						self.second_table[elem.row()][elem.column()] = str(elem.text())
			
			self.table_create(self.tableWidget_2, self.second_table)
			self.calculate_money()

	#third step calculate cost
	def calculation_pay(self):
		for i in range(len(self.decode[0])):
			self.decode[3][i] = round(self.pay['payme'] * self.coeff[self.decode[0][i]]/110.2 * 171.2 * self.decode[1][i]/self.decode[2][i],2)
			self.decode[4][i] = round(self.decode[3][i] * self.decode[2][i] / (171.2 * self.decode[1][i]), 2)
		self.decode[4][self.user] = round(self.decode[4][self.user] + self.pay['payme'] - sum(self.decode[4]), 2)
		self.decode[3][self.user] = round(self.decode[4][self.user]* 171.2 * self.decode[1][self.user] / self.decode[2][self.user], 2)

	#fourth step
	def calculation_fill(self):
		self.decode_buf = self.decode[3]+self.decode [4]
		self.decode_buf.append(self.pay['payme'])
		i = 0
		
		for key in sorted(list(self.pay_calc.keys())):
			self.pay_calc[key] = self.decode_buf[i]
			i += 1

	def change(self): #изменение ячейки
		query = self.sender().text()
		if len(query) >= 3:
			if self.sender().objectName()[-1] == '2':
				self.comboBox.clear()
				for key in self.prices.keys():
					if query.lower() in key.lower():
						self.comboBox.addItem(key)
				if len(self.comboBox.currentText()) != 0:
					self.first_table = self.prices[self.comboBox.currentText()]
					self.table_create(self.tableWidget, self.first_table)
			else:
				self.search_value(query)

	def calculate_money(self):
		if len(str(self.sale.text())) != 0:
			sum = 0
			for elem in self.second_table:
				sum += elem[-1]*elem[-2]
			self.sum_raw = round(sum*(100-int(self.sale.text()))/100, 2)
			self.sum_full = round(self.sum_raw*1.2*self.koeff/10000, 0)
			self.label_13.setText(str(self.sum_raw))
			self.label_11.setText(str(self.sum_full))

	#second_step count money
	def count_money(self):
		self.pay['allmoney'] = round(float(self.label_11.text()), 2)
		self.pay['ndssteal'] = round(float(self.pay['allmoney'])/6, 2)
		self.pay['withoutndssteal']= round(self.pay['allmoney'] - self.pay['ndssteal'], 2)
		self.pay['profit'] = round(float(self.pay['withoutndssteal'])/11, 2)
		self.pay['payandsteal'] = round(float(self.pay['withoutndssteal'])*10/11, 2)
		self.pay['payme'] = round(float(self.pay['payandsteal'])/1.691, 2)
		self.pay['stealins'] = round(float(self.pay['payme'])*0.001, 2)
		self.pay['stealsoc'] = round(float(self.pay['payme'])*0.34, 2)
		self.pay['stealnakl'] = round(self.pay['payandsteal'] - self.pay['stealins'] - self.pay['stealsoc'] - self.pay['payme'], 2)

	def combochange(self):
		if len(self.comboBox.currentText()) > 0:
			self.first_table = self.prices[self.comboBox.currentText()]
			self.table_create(self.tableWidget, self.first_table)

	#decoding
	#tab = [post, rate, pay, hours, full]
	def decode_pay(self):
		for row in self.doc.tables[self.tab_calc].rows[2:-1]:
			#print('\''+row.cells[3].text+'\'' + ' : '+'\'\'')
			self.decode[0].append(row.cells[0].text)
			self.decode[1].append(float(row.cells[1].text.replace(',', '.')))
			self.decode[2].append(float(row.cells[-3].text.replace(',', '.')))
			self.decode[3].append(0)
			self.decode[4].append(0)


	def delete_row(self):
		if self.tableWidget.selectedItems():
			for elem in self.tableWidget_2.selectedItems():
				self.second_table.pop(elem.row())
				self.table_create(self.tableWidget_2, self.second_table)
				self.calculate_money()

	def export_to_doc(self):
		if self.checkBox.isChecked():
			self.doc = docx.Document('dops.docx')
			self.tab_calc = 14
			self.tab_work = 12
			self.sec_sklon = 6
			self.timestamp['dayfr'] = str(self.textfield_14.text())
			self.timestamp['month'] = str(self.textfield_15.text())
			self.timestamp['year'] = str(self.textfield_16.text())
			self.timestamp['daytll'] = str(self.textfield_17.text())
			self.timestamp['mnthtill'] = str(self.textfield_19.text())
			self.timestamp['yeatill'] = str(self.textfield_18.text())
		else:
			self.doc = docx.Document('default.docx')
			self.tab_calc = 16
			self.tab_work = 14
			self.sec_sklon = 7
		self.customer = {'company':str(self.textfield_1.text()), 'comdir':str(self.textfield_2.text()), 'comdol':str(self.textfield_12.text()),\
			'compdoc':str(self.textfield_3.text()), 'comaddr':str(self.textfield_4.text()), 'combank':str(self.textfield_5.text()),\
			 'comunp':str(self.textfield_6.text()), 'comletter':str(self.textfield_7.text()), 'theme':str(self.textfield_8.text()),\
			  'themadd': str(self.plainTextEdit.toPlainText())}
		
		self.fontstyles = self.doc.styles
		self.user = self.comboBox_2.currentIndex()+2
		self.int_units = ((u'рубль', u'рубля', u'рублей'), 'm')
		self.exp_units = ((u'копейка', u'копейки', u'копеек'), 'f')
		self.decode  = [[], [], [], [], []]
		check = self.check_export()
		if check == 1:
			with open('users.txt', 'r') as bntu_val:
				for line in bntu_val:
					self.bntu[line.split('%')[0]] = line.split('%')[1][:-1]
			with open('workers.txt','r') as f:

				for line in f.readlines():
					for elem in (line[:-2].split('%')):
						if self.comboBox_3.currentText() in elem:
							self.bntu['nipi'] = elem.split('@')[0]
							self.bntu['nipidoc'] = elem.split('@')[1]
							self.bntu['nipidol'] = elem.split('@')[2]
					break
			self.bntu['boss'] = self.comboBox_2.currentText()
			self.bntu['chef'] = self.chefbox.text()
			self.bntu['zavpes'] = self.zavpes.text()
			for row in self.doc.tables[self.tab_calc].rows[2:]:
				for j in [-1, -2]:
					if len(row.cells[j].text)!= 0:
						self.pay_calc[row.cells[j].text] = ''
			self.decode_pay()
			self.count_money()
			self.calculation_pay()
			self.num_to_text()
			self.calculation_fill()
			#big_Dictionary
			self.zam_yes = False
			self.final = {}
			self.final.update(self.customer)
			self.final.update(self.pay)
			self.final.update(self.w_pay)
			self.final.update(self.timestamp)
			self.final.update(self.bntu)
			self.final.update(self.pay_calc)
			if self.checkBox.isChecked():
				self.final['dogovor'] = self.textfield_10.text()
				self.final['etap'] = self.textfield_9.text()
				self.final['dopsn'] = str(int(self.final['etap'])-1)
				allmon = float(self.textfield_11.text())+int(self.textfield_13.text())/100+float(self.label_11.text())
				self.final['sumpred'] = f"{int(allmon)} руб. {int((allmon*100)%100)} коп.".replace(' 0 ', ' 00 ')
				self.final['wstumpred'] = decimal2text(str(allmon),int_units = self.int_units ,exp_units =self.exp_units ).capitalize()
			for i in range(len(self.doc.paragraphs)):
				self.fill_text(self.doc.paragraphs[i], i)
			self.read_calc()
			self.fill_calc()
			file_name = QFileDialog.getSaveFileName(self, 'Save file', '/', filter='*.docx')[0]
			if file_name:
				self.doc.save(file_name)
				self.warning_wrong('Сохранение успешно!', 'Успех')
		elif check == 0:
			self.warning_wrong('Вы должны выдать заключение, добавьте!', 'Внимание!')
		elif check == 2:
			self.warning_wrong('Заполните все поля!', 'Внимание!')	

	def check_export(self):
		ans = 0
		for elem in self.second_table:
			if 'Выдача заключений' in elem[0]:
				ans = 1
		for key in list(self.customer.keys())[:-1]:
			if len(str(self.customer[key])) == 0:
				ans = 2

		return ans


	def init_base(self):
		with open('data.pi', 'rb') as f:
			self.prices = load(f)
		for key in self.prices.keys():
			self.comboBox.addItem(key)

	def num_to_text(self):
		self.w_pay['dgwithoutndssteal'] = f"{int(self.pay['withoutndssteal'])} руб. {int(self.pay['withoutndssteal']*100%100)} коп.".replace(' 0 ', ' 00 ')
		self.w_pay['dgndssteal'] = f"{int(self.pay['ndssteal'])} руб. {int(self.pay['ndssteal']*100%100)} коп.".replace(' 0 ', ' 00 ')
		self.w_pay['dgallmoney'] = f"{int(self.pay['allmoney'])} руб. {int(self.pay['allmoney']*100%100)} коп.".replace(' 0 ', ' 00 ')
		self.w_pay['wwithoutndssteal'] = decimal2text(str(self.pay['withoutndssteal']),int_units = self.int_units ,exp_units =self.exp_units ).capitalize()
		self.w_pay['wndssteal'] = decimal2text(str(self.pay['ndssteal']),int_units =self.int_units , exp_units =self.exp_units ).capitalize()
		self.w_pay['wallmoney'] = decimal2text(str(self.pay['allmoney']),int_units =self.int_units , exp_units =self.exp_units ).capitalize()

	#sixth step fill paragraph
	def fill_text(self, element,i):
		for run in element.runs:
			if (run.text.replace(' ', '')) in self.final.keys():
				if (i in [2, self.sec_sklon]) and (run.text.replace(' ', '') in ['comdol','comdir', 'nipidol', 'nipi']) and self.dops:
					reptext = str(self.sklon(self.final[run.text.replace(' ', '')]))
				else:
					reptext = str(self.final[run.text.replace(' ', '')])
				run.text=run.text.replace(run.text, reptext)
				#if 'заместитель' in reptext:
					#self.zam_yes = True
			#elif self.zam_yes and ('директор' in run.text):
				#run.text=run.text.replace(run.text, 'директора ')
				#self.zam_yes = False

	def sklon(self, phrase):
		reptext = ''
		for elem in phrase.split(' '):
			if elem[-1] in 'бвгджзклмнпрстфцчшщ':
				reptext += str(elem)+'а'+' '
			elif elem[-1] == 'ь':
				reptext += str(elem)[:-1]+'я'+' '
			elif elem[-1] == 'й':
				if elem[-2] == 'ы':
					reptext += str(elem)[:-2]+'ого'+' '
				else:
					reptext += str(elem)[:-2]+'его'+' '
			else:
				reptext += elem+' '
		return reptext[:-1]

	def fill_users(self):
		with open('workers.txt','r') as f:
			i=0
			for line in f.readlines():
				for elem in (line[:-2].split('%')):
					if i == 0:
						self.comboBox_3.addItem(elem.split('@')[0])
					else:
						self.comboBox_2.addItem(elem.split('@')[0])
				i += 1

	def fill_calc(self):
		for elem in self.second_table[:-1]:
			cells = self.doc.tables[self.tab_work].add_row().cells
			for i in range(2):
				cells[i].text = str(elem[i])


	#seventh step fill tables
	def read_calc(self):
		for i in range(len(self.doc.tables)):
			if i == 2:
				self.dops = False
			else:
				self.dops = True
			for j in range(len(self.doc.tables[i].rows)):
				for k in range(len(self.doc.tables[i].rows[0].cells)):
					cell = self.doc.tables[i].cell(j, k)
					for L in range(len(cell.paragraphs)):
						self.fill_text(cell.paragraphs[L], i)

	def Save_file(self): #сохранить файл
		fname = QFileDialog.getSaveFileName(self, 'Save file', '/', filter='*.txt')[0]
		if fname:
			with open(fname, 'w') as f:
				for elem in self.fields:
					f.write(elem.text().replace('\n', ' '))
					f.write('\n')					

	def Open_file(self): #открыть файл
		fname = QFileDialog.getOpenFileName(self, 'Open file', '/', filter='*.txt')[0]
		if fname:
			with open(fname, 'r') as f:
				i = 0
				for line in f.readlines():
					if i>=len(self.fields):
						self.plainTextEdit.setPlainText(str(line)[:-1])
						break
					self.fields[i].setText(str(line)[:-1])
					i+=1

	def search_value(self, query):
		self.first_table = []
		for elem in self.prices.values():
			for word in elem:
				if query.lower() in word[0].lower():
					self.first_table.append(word)
		self.table_create(self.tableWidget, self.first_table)

	def table_create(self, table, tab_files): #создание таблицы
		self.table_show = 1 #не учитываем изменения ячеек
		table.clearContents()#очищаем таблицу
		table.setRowCount(len(tab_files))        # Устанавливаем строки в таблице
		if len(tab_files) != 0:
			for i in range(len(tab_files[0])):
				count = 0
				for j in tab_files:
					table.setItem(count, i, QTableWidgetItem(str(j[i])))
					count += 1
		table.resizeRowsToContents()
		self.table_show = 0

	def warning_wrong(self, warn, mis): #Ошибка
		msg = QMessageBox()
		msg.setIcon(QMessageBox.Warning)
		msg.setText(mis)
		msg.setInformativeText(warn)
		msg.setWindowTitle(mis)
		msg.setStandardButtons(QMessageBox.Ok)
		msg.exec_()

if __name__ == "__main__":
	import sys

	app = QApplication(sys.argv)
	w = Ui()
	w.show()  # show window
	sys.exit(app.exec_())

 